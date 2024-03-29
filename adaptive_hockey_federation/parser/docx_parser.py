import re
from datetime import date
from typing import Optional

import docx  # type: ignore

from adaptive_hockey_federation.parser.user_card import BaseUserInfo

NAME = '[И|и][М|м][Я|я]|Ф.И.О.'
SURNAME = '[Ф|ф][А|а][М|м][И|и][Л|л][И|и][Я|я]|Ф.И.О.'
PATRONYMIC = '[О|о][Т|т]?[Ч|ч][Е|е][С|с][Т|т][В|в][О|о]|Ф.И.О.'
DATE_OF_BIRTH = '[Д|д][А|а][Т|т][А|а] [Р|р][О|о].+'
TEAM = '[К|к][О|о][М|м][А|а][Н|н][Д|д][А|а]'
PLAYER_NUMBER = '[И|и][Г|г][Р|р][О|о][В|в][О|о][Й|й]'
POSITION = '[П|п][О|о][З|з][И|и][Ц|ц][И|и][Я|я]|Должность'
NUMERIC_STATUS = '[Ч|ч].+[С|с][Т|т].+'
PLAYER_CLASS = '[К|к][Л|л][А|а][С|с][С|с]'
PASSPORT = '[П|а][С|с][П|п][О|о][Р|р][Т|т]'
ASSISTENT = ['А', 'а', '(А)', '(а)', 'Ассистент', 'ассистент']
CAPTAIN = ['К', 'к', '(К)', '(к)', 'Капитан', 'капитан']
DISCIPLINE_LEVEL = 'без ограничений'


def read_file_columns(file: docx) -> list[docx]:
    """Функция находит таблицы в файле и возвращает список объектов
    docx с данными каждого столбца.
    """
    return [
        column
        for table in file.tables
        for index, column in enumerate(table.columns)
    ]


def read_file_text(file: docx) -> list[str]:
    """Функция находит текстовые данные в файле и возвращает список объектов
    docx с найденными данными.
    """
    return [
        run.text
        for paragraph in file.paragraphs
        for run in paragraph.runs
    ]


def get_counter_for_columns_parser(
        columns: list[docx]
) -> int:
    count = 0
    for column in columns:
        for index, cell in enumerate(column.cells):
            if re.search(r'п/п', cell.text):
                for cell in column.cells[index + 1:]:
                    if cell.text and len(cell.text) < 4:
                        count += 1
                    else:
                        break
            else:
                if count > 0:
                    break
    return count


def columns_parser(
        columns: list[docx],
        regular_expression: str,
) -> list[Optional[str]]:
    """Функция находит столбец по названию и списком выводит содержимое
    каждой ячейки этого столбца.
    """
    output = [
        text if text
        else None
        for column in columns
        if re.search(
            regular_expression,
            list(cell.text for cell in column.cells)[0]
        )
        for text in list(cell.text for cell in column.cells)[1:]
    ]
    if not output:
        count = get_counter_for_columns_parser(columns)
        for column in columns:
            for index, cell in enumerate(column.cells):
                if re.search(regular_expression, cell.text):
                    for cell in column.cells[index + 1:index + 1 + count]:
                        output.append(cell.text)
    return output


def find_names(columns: list[docx], regular_expression: str) -> list[str]:
    """Функция парсит в искомом столбце имена. Опирается на шаблон ФИО
    (имя идет после фамилии на втором месте).
    """
    names_list = columns_parser(columns, regular_expression)
    return [
        name.split()[1].rstrip()
        for name in names_list
        if name
    ]


def find_surnames(columns: list[docx], regular_expression: str) -> list[str]:
    """Функция парсит в искомом столбце фамилии. Опирается на шаблон ФИО
    (фамилия идет на первом месте).
    """
    surnames_list = columns_parser(columns, regular_expression)
    return [
        surname.split()[0].rstrip()
        for surname in surnames_list
        if surname
    ]


def find_patronymics(
        columns: list[docx],
        regular_expression: str,
) -> list[str]:
    """Функция парсит в искомом столбце отчества. Опирается на шаблон ФИО
    (отчество идет на последнем месте).
    """
    patronymics_list = columns_parser(columns, regular_expression)
    return [
        patronymic.replace('/', ' ').split()[2].rstrip().rstrip(',')
        if patronymic and len(patronymic.split()) > 2
        else 'Отчество отсутствует'
        for patronymic in patronymics_list
    ]


def find_dates_of_birth(
        columns: list[docx],
        regular_expression: str,
) -> list[date]:
    """Функция парсит в искомом столбце дату рождения
    и опирается на шаблон дд.мм.гггг.
    """
    dates_of_birth_list = columns_parser(columns, regular_expression)
    dates_of_birth_list_clear = []
    for date_of_birth in dates_of_birth_list:
        if date_of_birth:
            try:
                for day, month, year in [
                    re.sub(r'\D', ' ', date_of_birth).split()
                ]:
                    if len(year) == 2:
                        if int(year) > 23:
                            year = '19' + year
                        else:
                            year = '20' + year
                    dates_of_birth_list_clear.append(
                        date(int(year), int(month), int(day))
                    )
            except ValueError or IndexError:  # type: ignore
                dates_of_birth_list_clear.append(date(1900, 1, 1))
        else:
            dates_of_birth_list_clear.append(date(1900, 1, 1))

    return dates_of_birth_list_clear


def find_team(
        text: list[str],
        columns: list[docx],
        regular_expression: str,
) -> str:
    """Функция парсит название команды.
    """
    text_clear = ' '.join(text)
    text_clear = re.sub(
        r'\W+|_+|ХК|СХК|ДЮСХК|Хоккейный клуб|по незрячему хоккею'
        '|по специальному хоккею|Спец хоккей|по специальному|по следж-хоккею',
        ' ',
        text_clear
    ).split()  # type: ignore
    try:
        return [
            'Молния Прикамья'
            if text_clear[index + 2] == 'Прикамья'
            else 'Ак Барс'
            if text_clear[index + 1] == 'Ак'
            else 'Снежные Барсы'
            if text_clear[index + 1] == 'Снежные'
            else 'Хоккей Для Детей'
            if text_clear[index + 1] == 'Хоккей'
            else 'Дети-Икс'
            if text_clear[index + 1] == 'Дети'
            else 'СКА-Стрела'
            if text_clear[index + 1] == 'СКА'
            else 'Сборная Новосибирской области'
            if text_clear[index + 2] == 'Новосибирской'
            else 'Атал'
            if text_clear[index + 3] == 'Атал'
            else 'Крылья Мечты'
            if text_clear[index + 2] == 'мечты'
            else 'Огни Магнитки'
            if text_clear[index + 1] == 'Огни'
            else 'Энергия Жизни Краснодар'
            if text_clear[index + 3] == 'Краснодар'
            else 'Энергия Жизни Сочи'
            if text_clear[index + 4] == 'Сочи'
            else 'Динамо-Москва'
            if text_clear[index + 1] == 'Динамо'
            else 'Крылья Советов'
            if text_clear[index + 2] == 'Советов'
            else 'Красная Ракета'
            if text_clear[index + 2] == 'Ракета'
            else 'Красная Молния'
            if text_clear[index + 2] == 'молния'
            else 'Сахалинские Львята'
            if text_clear[index + 1] == 'Сахалинские'
            else 'Мамонтята Югры'
            if text_clear[index + 1] == 'Мамонтята'
            else 'Уральские Волки'
            if text_clear[index + 1] == 'Уральские'
            else 'Нет названия команды'
            if text_clear[index + 1] == 'Всего'
            else text_clear[index + 1].capitalize()
            for index, txt in enumerate(text_clear)
            if re.search(regular_expression, txt)
        ][0]
    except IndexError:
        for column in columns:
            for cell in column.cells:
                if re.search(regular_expression, cell.text):
                    txt = re.sub(r'\W', ' ', cell.text)
                    return txt.split()[1].capitalize()

        return 'Название команды не найдено'


def find_players_number(
        columns: list[docx],
        regular_expression: str,
) -> list[int]:
    """Функция парсит в искомом столбце номер игрока.
    """
    players_number_list = columns_parser(columns, regular_expression)
    players_number_list_clear = []
    for player_number in players_number_list:
        if player_number:
            try:
                players_number_list_clear.append(
                    int(re.sub(r'\D', '', player_number)[:2])
                )
            except ValueError:
                players_number_list_clear.append(0)
        else:
            players_number_list_clear.append(0)

    return players_number_list_clear


def find_positions(columns: list[docx], regular_expression: str) -> list[str]:
    """Функция парсит в искомом столбце позицию игрока на поле.
    """
    positions_list = columns_parser(columns, regular_expression)
    return [
        'нападающий'
        if re.search(
            r'^н|^Н|^H|^Нп|^нл|^нп|^цн|^лн|^Нап|^№|^А,|^К,',
            position.lstrip()
        )
        else 'защитник'
        if re.search(r'^з|^З|^Зщ|^Защ', position.lstrip())
        else 'вратарь'
        if re.search(r'^Вр|^В|^вр', position.lstrip())
        else 'Позиция записана неверно'
        if not re.sub(r'\n|\(.+|\d', '', position)
        else re.sub(
            r'\n|\(.+|\d|Капитан',
            '',
            position
        ).lower().rstrip().replace(',', '').lstrip()
        for position in positions_list
        if position
    ]


def find_numeric_statuses(file: docx) -> list[list[str]]:
    numeric_statuses_list = []
    for table in file.tables:
        for row in table.rows:
            txt = row.cells[1].text.title()
            txt = re.sub(r'\W|Коляс.+|Здоровый', ' ', txt)
            if len(txt.split()) <= 4:
                try:
                    numeric_status = row.cells[4].text
                    numeric_status = re.sub(r'\D', '', numeric_status)
                    if numeric_status:
                        if len(txt.split()) == 2:
                            txt += ' Отчество отсутствует'
                        numeric_statuses_list.append(
                            txt.split()[:3] + [numeric_status]
                        )
                except IndexError:
                    pass

    return numeric_statuses_list


def find_passport(columns: list[docx], regular_expression: str) -> list[str]:
    """Функция парсит в искомом столбце ПД.
    """
    identity_list = columns_parser(columns, regular_expression)
    return [
        ' '.join(identity_data.split())
        for identity_data in identity_list
        if identity_data
    ]


def find_players_is_captain(
        columns: list[docx],
        regular_expression: str
) -> list[bool]:
    """Функция парсит в искомом столбце капитанов.
    """
    is_captain_list = []
    for is_captain in columns_parser(columns, regular_expression):
        for i in CAPTAIN:
            if is_captain and i in is_captain.strip():
                try:
                    is_captain_list.append(True)
                except ValueError:
                    is_captain_list.append(False)
            else:
                is_captain_list.append(False)
    return is_captain_list


def find_players_is_assistant(
        columns: list[docx],
        regular_expression: str,
) -> list[bool]:
    """Функция парсит в искомом столбце асситсента.
    """
    is_assistant_list = []
    for is_assistant in columns_parser(columns, regular_expression):
        for i in ASSISTENT:
            if is_assistant and i in is_assistant.strip():
                try:
                    is_assistant_list.append(True)
                except ValueError:
                    is_assistant_list.append(False)
            else:
                is_assistant_list.append(False)
    return is_assistant_list


def find_discipline_level(
        columns: list[docx],
        regular_expression: str,
) -> list[str]:
    """Функция парсит в искомом столбце класс/статус.
    """
    discipline_level_list = []
    for discipline_level in columns_parser(columns, regular_expression):
        if discipline_level:
            try:
                discipline_level = (discipline_level.replace(
                    'Класс ', '').replace('класс ', '')
                )
                discipline_level = re.sub(
                    r'без ограничений|Не имеет ограничений по здоровью'
                    '|Без ограничений по здоровью'
                    '|4Без ограничений'
                    '|Без ограничений\n4|без ограничений по здоровью 2'
                    '|игрок без ограничений по здоровью'
                    '|(без ограничений по здоровью)  3'
                    '|без ограничений 2|Не имеет ограничений по здоровью',
                    'б\\к', discipline_level)
                if discipline_level != 'б\\к':
                    discipline_level = (
                        discipline_level.replace('А', 'A').replace(
                            'С', 'C').replace('Б', 'B').replace('б', 'B')
                    )
                discipline_level_list.append(discipline_level)
            except ValueError:
                discipline_level_list.append('')
        else:
            discipline_level_list.append('')
    return discipline_level_list


def numeric_status_check(
        name: str,
        surname: str,
        patronymics: str,
        statuses: list[list[str]],
) -> Optional[int]:
    for status in statuses:
        if surname == status[0]:
            if name == status[1]:
                if patronymics.split()[0] == status[2]:
                    return int(status[3])

    return None


def length_list(name: list, len_name: int) -> None:
    if len(name) != len_name:
        for _ in range(len_name - len(name)):
            name.append(None)
    return None


def docx_parser(
        path: str,
        numeric_statuses: list[list[str]]
) -> list[BaseUserInfo]:
    """Функция собирает все данные об игроке
    и передает их в dataclass.
    """
    file = docx.Document(path)
    columns_from_file = read_file_columns(file)
    text_from_file = read_file_text(file)
    names = find_names(columns_from_file, NAME)
    surnames = find_surnames(columns_from_file, SURNAME)
    patronymics = find_patronymics(columns_from_file, PATRONYMIC)
    dates_of_birth = find_dates_of_birth(
        columns_from_file,
        DATE_OF_BIRTH,
    )
    team = find_team(text_from_file, columns_from_file, TEAM)
    players_number = find_players_number(columns_from_file, PLAYER_NUMBER)
    positions = find_positions(columns_from_file, POSITION)
    passport = find_passport(columns_from_file, PASSPORT)
    is_assistents = find_players_is_assistant(columns_from_file, PLAYER_NUMBER)
    is_assistents_alt = find_players_is_assistant(columns_from_file, POSITION)
    is_captains = find_players_is_captain(columns_from_file, PLAYER_NUMBER)
    is_captains_alt = find_players_is_captain(columns_from_file, POSITION)
    classification = find_discipline_level(columns_from_file, DISCIPLINE_LEVEL)

    length_list(players_number, len(names))
    length_list(is_assistents, len(names))
    length_list(is_captains, len(names))
    length_list(classification, len(names))
    length_list(passport, len(names))
    length_list(positions, len(names))

    return [
        BaseUserInfo(
            name=names[index],
            surname=surnames[index],
            date_of_birth=dates_of_birth[index],
            team=team,
            player_number=players_number[index],
            position=positions[index],
            numeric_status=numeric_status_check(
                names[index],
                surnames[index],
                patronymics[index],
                numeric_statuses,
            ),
            patronymic=patronymics[index],
            passport=passport[index],
            is_assistant=(is_assistents[index] if is_assistents[index]
                          else is_assistents_alt[index]),
            is_captain=(is_captains[index] if is_captains[index]
                        else is_captains_alt[index]),
            classification=classification[index]
        ).__dict__
        for index in range(len(names))
    ]
